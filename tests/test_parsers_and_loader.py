from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from localrag.ingestion import loader as ingestion_loader
from localrag.ingestion.loader import list_supported_files, parse_file
from localrag.ingestion.parsers.docx import parse_docx
from localrag.ingestion.parsers.markdown import parse_markdown
from localrag.ingestion.parsers.text import parse_text


def test_parse_markdown_front_matter(tmp_path: Path) -> None:
    path = tmp_path / "doc.md"
    path.write_text(
        "---\ntitle: x\n---\n\nhello\nworld\n",
        encoding="utf-8",
    )
    out = parse_markdown(path)
    assert out == "hello\nworld"


def test_parse_text_empty_bytes_returns_empty_string(tmp_path: Path) -> None:
    path = tmp_path / "a.txt"
    path.write_bytes(b"")
    out = parse_text(path)
    assert out == ""


def test_parse_docx_parses_paragraphs(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    path = tmp_path / "a.docx"
    doc.save(path)

    out = parse_docx(path)
    assert out == "Hello\nWorld"


def test_loader_list_supported_files_recursive(tmp_path: Path) -> None:
    root = tmp_path / "root"
    (root / "nested").mkdir(parents=True)

    (root / "a.md").write_text("x", encoding="utf-8")
    (root / "nested" / "b.md").write_text("y", encoding="utf-8")
    (root / "nested" / "c.txt").write_text("z", encoding="utf-8")

    files = list_supported_files(root, recursive=True)

    names = sorted(p.name for p in files)
    assert names == ["a.md", "b.md", "c.txt"]


def test_loader_list_supported_files_when_path_is_file(tmp_path: Path) -> None:
    md = tmp_path / "a.md"
    md.write_text("x", encoding="utf-8")

    out = list_supported_files(md, recursive=False)
    assert out == [md]


def test_loader_parse_file_dispatches_extensions(tmp_path: Path) -> None:
    md = tmp_path / "a.md"
    md.write_text("---\nmeta\n---\nhello", encoding="utf-8")

    txt = tmp_path / "b.txt"
    txt.write_bytes(b"hello")

    py = tmp_path / "c.py"
    py.write_text("print('x')", encoding="utf-8")

    doc = Document()
    doc.add_paragraph("p1")
    doc.add_paragraph("p2")
    docx = tmp_path / "d.docx"
    doc.save(docx)

    assert parse_file(md).endswith("hello")
    assert parse_file(txt) == "hello"
    assert parse_file(py) == "print('x')"
    assert parse_file(docx) == parse_docx(docx)


def test_loader_parse_file_dispatches_pdf(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    # Avoid real PDF parsing in unit tests: we only want to test dispatch.
    monkeypatch.setattr(ingestion_loader, "parse_pdf", lambda _: "PDF ok")

    pdf_path = tmp_path / "a.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%fake\n")

    assert parse_file(pdf_path) == "PDF ok"
